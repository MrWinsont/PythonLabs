from docx import Document
def AddPicture(image_path, docx_path):
    doc = Document(docx_path)
    doc.add_picture(image_path)
    doc.add_paragraph('изображение успешно сохранено')
    doc.save('file.docx')

AddPicture('image.jpg', 'file.docx')